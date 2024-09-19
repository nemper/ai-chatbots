# in myfunc.embeddings.py
import ast
import chardet
import cohere
import json
import logging
import os
import streamlit as st
import io
import re
import nltk
nltk.download('punkt')
from datetime import datetime
from bs4 import BeautifulSoup

from openai import OpenAI
from pinecone import Pinecone
from pinecone_text.sparse import BM25Encoder
from time import sleep
from tqdm.auto import tqdm
from uuid import uuid4
import pandas as pd

from langchain.chains import GraphQAChain
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.graphs.networkx_graph import NetworkxEntityGraph
from langchain.output_parsers import ResponseSchema, StructuredOutputParser
from langchain.prompts import PromptTemplate
from langchain.retrievers.multi_query import MultiQueryRetriever
from langchain_text_splitters import RecursiveCharacterTextSplitter, HTMLHeaderTextSplitter
from langchain_community.document_loaders import PDFMinerPDFasHTMLLoader
from langchain_community.document_transformers import LongContextReorder
from langchain_community.retrievers import PineconeHybridSearchRetriever
from langchain_openai import ChatOpenAI, OpenAIEmbeddings

from semantic_router.encoders import OpenAIEncoder
from semantic_router.splitters import RollingWindowSplitter
# from semantic_router.utils.logger import logger

from myfunc.mojafunkcija import st_style, pinecone_stats
from myfunc.retrievers import HybridQueryProcessor, PineconeUtility, SelfQueryPositive, TextProcessing
from myfunc.various_tools import get_structured_decision_from_model, positive_calendly, web_search_process, scrape_webpage_text, hyde_rag
from myfunc.mssql import work_prompts

import markdown
import pypandoc
from langchain.docstore.document import Document

mprompts = work_prompts()

st_style()
client=OpenAI()
text_processor = TextProcessing(gpt_client=client)
pinecone_utility = PineconeUtility()



# in myfunc.embeddings.py
class DocumentConverter:
    """
    A class for converting documents between different formats and extracting structured data.

    This class provides methods to convert documents from Markdown, DOCX, and PDF formats to HTML,
    and to split HTML content based on headers. It also processes PDF documents to extract
    structured data based on font size, identifying headings and associated content.

    Methods:
    - conv_md(tekst): Converts a Markdown file to HTML.
    - conv_docx(tekst): Converts a DOCX file to HTML.
    - split_on_headers(html_string): Splits an HTML document into sections based on specified header levels.
    - conv_pdf(tekst): Converts a PDF file into a list of structured data snippets, each representing a section of the document.

    The class is designed to facilitate the handling of different document formats and to streamline the
    extraction of meaningful and structured content from these documents.
    """
    def __init__(self):
        self.headers_to_split_on = [("h2", "heading")]
        # Additional initialization can go here

    def conv_md(self, tekst):
        with open(tekst, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
        return markdown.markdown(markdown_text)

    def conv_docx(self, tekst):
        return pypandoc.convert_file(tekst, 'html')

    def split_on_headers(self, html_string):
        html_splitter = HTMLHeaderTextSplitter(headers_to_split_on=self.headers_to_split_on)
        return html_splitter.split_text(html_string)
    

    def conv_csv(self, file_path):
        df = pd.read_csv(file_path)
        documents = []
        for _, row in df.iterrows():
            content = " ".join([f"{col}: {row[col]}" for col in df.columns])
            metadata = {col: row[col] for col in df.columns}
            documents.append(Document(page_content=content, metadata=metadata))
        return documents
    
        
    def conv_pdf(self, doc_path):  
        loader = PDFMinerPDFasHTMLLoader(doc_path)
        data = loader.load()[0]   # entire PDF is loaded as a single Document
        soup = BeautifulSoup(data.page_content,'html.parser')
        content = soup.find_all('div')
        cur_fs = None
        cur_text = ''
        snippets = []   # first collect all snippets that have the same font size
        for c in content:
            sp = c.find('span')
            if not sp:
                continue
            st = sp.get('style')
            if not st:
                continue
            fs = re.findall('font-size:(\d+)px',st)
            if not fs:
                continue
            fs = int(fs[0])
            if not cur_fs:
                cur_fs = fs
            if fs == cur_fs:
                cur_text += c.text
            else:
                snippets.append((cur_text,cur_fs))
                cur_fs = fs
                cur_text = c.text
        snippets.append((cur_text,cur_fs))
        # Note: The above logic is very straightforward. One can also add more strategies such as removing duplicate snippets (as
        # headers/footers in a PDF appear on multiple pages so if we find duplicates it's safe to assume that it is redundant info)
        cur_idx = -1
        semantic_snippets = []
        # Assumption: headings have higher font size than their respective content
        for s in snippets:
            # if current snippet's font size > previous section's heading => it is a new heading
            if not semantic_snippets or s[1] > semantic_snippets[cur_idx].metadata['heading_font']:
                metadata={'heading':s[0], 'content_font': 0, 'heading_font': s[1]}
                metadata.update(data.metadata)
                semantic_snippets.append(Document(page_content='',metadata=metadata))
                cur_idx += 1
                continue

            # if current snippet's font size <= previous section's content => content belongs to the same section (one can also create
            # a tree like structure for sub sections if needed but that may require some more thinking and may be data specific)
            if not semantic_snippets[cur_idx].metadata['content_font'] or s[1] <= semantic_snippets[cur_idx].metadata['content_font']:
                semantic_snippets[cur_idx].page_content += s[0]
                semantic_snippets[cur_idx].metadata['content_font'] = max(s[1], semantic_snippets[cur_idx].metadata['content_font'])
                continue

            # if current snippet's font size > previous section's content but less than previous section's heading than also make a new
            # section (e.g. title of a PDF will have the highest font size but we don't want it to subsume all sections)
            metadata={'heading':s[0], 'content_font': 0, 'heading_font': s[1]}
            metadata.update(data.metadata)
            semantic_snippets.append(Document(page_content='',metadata=metadata))
            cur_idx += 1
        return semantic_snippets

from streamlit.runtime.uploaded_file_manager import UploadedFile


def is_binary(data):
    text_characters = bytearray({7, 8, 9, 10, 12, 13, 27} | set(range(0x20, 0x100)))
    return bool(data.translate(None, text_characters))

import PyPDF2
from docx import Document
from io import BytesIO
def handle_docx_file(raw_data):
    doc = Document(BytesIO(raw_data))
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return "\n".join(text)

def handle_pdf_file(raw_data):
    reader = PyPDF2.PdfReader(BytesIO(raw_data))
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text


import chardet
import re
import streamlit as st
from streamlit.runtime.uploaded_file_manager import UploadedFile


def read_uploaded_file(uploaded_file, text_delimiter=" "):
    if isinstance(uploaded_file, UploadedFile):
        # Streamlit's UploadedFile needs to be read directly
        raw_data = uploaded_file.read()

        if is_binary(raw_data):
            # Determine if it's a DOCX or PDF file by checking the file extension
            if uploaded_file.name.endswith(".docx"):
                data = handle_docx_file(raw_data)
            elif uploaded_file.name.endswith(".pdf"):
                data = handle_pdf_file(raw_data)
            else:
                st.write("Binary file detected but unsupported format for text extraction.")
                return None

            # Process the extracted text with the text delimiter
            if data is not None and text_delimiter:
                data = process_text_with_delimiter(data, text_delimiter)

            return data

        # Handle text files by detecting encoding
        result = chardet.detect(raw_data)
        encoding = result['encoding']
        try:
            data = raw_data.decode(encoding, errors="replace")  # Handle unknown characters gracefully
        except (UnicodeDecodeError, TypeError):
            data = raw_data.decode("windows-1252", errors="replace")  # Fallback to windows-1252

        # Process the decoded text with the text delimiter
        if text_delimiter:
            data = process_text_with_delimiter(data, text_delimiter)

    else:
        # Handle as a regular file path
        try:
            with open(uploaded_file, 'rb') as f:  # Read file as binary
                raw_data = f.read()
            
            if is_binary(raw_data):
                # Determine if it's a DOCX or PDF file
                if uploaded_file.endswith(".docx"):
                    data = handle_docx_file(raw_data)
                elif uploaded_file.endswith(".pdf"):
                    data = handle_pdf_file(raw_data)
                else:
                    st.write("Binary file detected but unsupported format for text extraction.")
                    return None

                # Process the extracted text with the text delimiter
                if data is not None and text_delimiter:
                    data = process_text_with_delimiter(data, text_delimiter)

                return data

            # Handle text files by detecting encoding
            result = chardet.detect(raw_data)
            encoding = result['encoding']
            try:
                data = raw_data.decode(encoding, errors="replace")
            except (UnicodeDecodeError, TypeError):
                data = raw_data.decode("windows-1252", errors="replace")  # Fallback to windows-1252

            # Process the decoded text with the text delimiter
            if text_delimiter:
                data = process_text_with_delimiter(data, text_delimiter)
        except (UnicodeDecodeError, TypeError):
            data = raw_data.decode("windows-1252", errors="replace")  # Fallback to windows-1252

    return data



def process_text_with_delimiter2(text, text_delimiter):
    if not text_delimiter:
        text_delimiter = "\n\n"

    # Do not replace the delimiter; keep it in the text
    # Proceed with other text cleaning tasks

    # Additional text cleaning
    text = text.replace("•", "")
    text = re.sub(r"(?<=\b\w) (?=\w\b)", "", text)

    return text


def process_text_with_delimiter(text, text_delimiter):
    """
    Processes the text by ensuring the specified text delimiter is handled correctly.

    Parameters:
    - text: The text content to process.
    - text_delimiter: The delimiter used to split or replace in the text.

    Returns:
    - The processed text.
    """
    if not text_delimiter:
        text_delimiter = "\n\n"

    # Ensure the text uses consistent newlines
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Additional text cleaning before splitting
    text = text.replace("•", "")
    text = re.sub(r"(?<=\b\w) (?=\w\b)", "", text)

    return text


def is_binary(data):
    # Placeholder for binary detection logic
    return b'\0' in data

def handle_docx_file(raw_data):
    # Placeholder for handling DOCX files
    return extract_text_from_docx(raw_data)

def handle_pdf_file(raw_data):
    # Placeholder for handling PDF files
    return extract_text_from_pdf(raw_data)



def read_uploaded_file2(uploaded_file):
    if isinstance(uploaded_file, UploadedFile):
        # Streamlit's UploadedFile needs to be read directly
        raw_data = uploaded_file.read()

        if is_binary(raw_data):
            # Determine if it's a DOCX or PDF file by checking the file extension
            if uploaded_file.name.endswith(".docx"):
                return handle_docx_file(raw_data)
            elif uploaded_file.name.endswith(".pdf"):
                return handle_pdf_file(raw_data)
            else:
                st.write("Binary file detected but unsupported format for text extraction.")
                return None

        # Handle text files by detecting encoding
        result = chardet.detect(raw_data)
        encoding = result['encoding']
        try:
            data = raw_data.decode(encoding, errors="replace")  # Handle unknown characters gracefully
        except (UnicodeDecodeError, TypeError):
            data = raw_data.decode("windows-1252", errors="replace")  # Fallback to windows-1252
    else:
        # Handle as a regular file path
        try:
            with open(uploaded_file, 'rb') as f:  # Read file as binary
                raw_data = f.read()
            
            if is_binary(raw_data):
                # Determine if it's a DOCX or PDF file
                if uploaded_file.endswith(".docx"):
                    return handle_docx_file(raw_data)
                elif uploaded_file.endswith(".pdf"):
                    return handle_pdf_file(raw_data)
                else:
                    st.write("Binary file detected but unsupported format for text extraction.")
                    return None

            # Handle text files by detecting encoding
            result = chardet.detect(raw_data)
            encoding = result['encoding']
            data = raw_data.decode(encoding, errors="replace")
        except (UnicodeDecodeError, TypeError):
            data = raw_data.decode("windows-1252", errors="replace")  # Fallback to windows-1252
    return data

import re
from langchain_text_splitters import RecursiveCharacterTextSplitter

def standard_chunks(dokum, chunk_size, chunk_overlap, sep="\n\n", keep=False):
    # Escape the separator to handle regex special characters
    escaped_sep = re.escape(sep)
    
    text_splitter = RecursiveCharacterTextSplitter(
        separators=[escaped_sep, "\n\n", "\n", " ", ""],
        keep_separator=keep,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        is_separator_regex=True,  # Set to True to use regex
    )
    
    # Read the content of the uploaded file
    data = read_uploaded_file(dokum)
    
    # Get a serializable file name or other string identifier
    source_identifier = dokum.name if isinstance(dokum, UploadedFile) else str(dokum)
    
    # Create documents by splitting the text
    texts = text_splitter.create_documents([data], metadatas=[{"source": source_identifier}])
    
    # Proceed with creating JSON as before
    output_json_list = []
    current_date = datetime.now()
    date_string = current_date.strftime('%Y%m%d')
    
    for i, document in enumerate(texts, start=1):
        output_dict = {
            "id": str(uuid4()),
            "chunk": i,
            "text": document.page_content.strip(),
            "source": document.metadata.get("source", ""),
            "date": int(date_string),
        }
        output_json_list.append(output_dict)
    
    json_string = (
        "["
        + ",\n".join(
            json.dumps(d, ensure_ascii=False) for d in output_json_list
        )
        + "]"
    )
    return json_string


# ovo funkcija radi
def semantic_chunks(sadrzaj, source):
    '''
    Ulaz je Document object iz loadera i naziv source fajla
    Izlaz je JSON string za embedding
    Ima opcije za ploting i statistike, dobro za testiranje optimalnih parametara
    '''
    encoder = OpenAIEncoder(name="text-embedding-3-large", dimensions = 3072)
    # encoder.score_threshold = 0.2 ako se postavi dinamic_treshold = False
    # logger.setLevel("WARNING")  # reduce logs from splitter
    splitter = RollingWindowSplitter(
        encoder=encoder,
        dynamic_threshold=True, # mora biti False ako se definise fixni
        min_split_tokens=100,
        max_split_tokens=500,
        window_size=5, # deafult
        plot_splits=False,  # set this to true to visualize chunking
        enable_statistics=False  # to print chunking stats
    )
    splits = splitter([sadrzaj[0].page_content])
    current_date = datetime.now()
    date_string = current_date.strftime('%Y%m%d')
    structured_data = []
    
    for i, s in enumerate(splits):  
            output_dict = {
                "id": str(uuid4()),
                "chunk": i,
                "text": s.content,
                "source": source,
                "date": int(date_string),
            }
            structured_data.append(output_dict)

                
    json_string = (
        "["
        + ",\n".join(
            json.dumps(d, ensure_ascii=False) for d in structured_data
        )
        + "]"
    )

   
    return json_string


def heading_chunks(uploaded_file):
    '''
    Ulaz je fajl iz st.upload_file
    Izlaz je JSON string za embedding
    '''
    # Instantiate the DocumentConverter class
    converter = DocumentConverter()

    current_date = datetime.now()
    date_string = current_date.strftime('%Y%m%d')
    structured_data = []

    # Handling file name and extension
    _, ext = os.path.splitext(uploaded_file.name)

    # Processing the document based on its extension
    if ext == ".pdf":
        document = converter.conv_pdf(uploaded_file.name)
    elif ext == ".docx":
        html = converter.conv_docx(uploaded_file.name)
        document = converter.split_on_headers(html)
    elif ext == ".md":
        html = converter.conv_md(uploaded_file.name)
        document = converter.split_on_headers(html)
    else:
        st.error("Only .md, .pdf, .csv and .docx files are supported.")
        return 

    
    i = 0
    for doc in document:
        i += 1
        title = doc.metadata.get('heading', "")
        content = doc.page_content
        output_dict = {
            "id": str(uuid4()),
            "chunk": i,
            "text": title + " > " + content,
            "heading": title,
            "source": uploaded_file.name,
            "date": int(date_string),
        }
        structured_data.append(output_dict)

    json_string = (
        "["
        + ",\n".join(
            json.dumps(d, ensure_ascii=False) for d in structured_data
        )
        + "]"
    )

   
    return json_string

def csv_chunks(uploaded_file):
    '''
    Ulaz je fajl iz st.upload_file
    Izlaz je JSON string za embedding
    '''
    _, ext = os.path.splitext(uploaded_file.name)
    if ext == ".csv" or ext == ".CSV":
        pass
    else:
        st.error("Only .csv is supported.")
        return 
    with st.spinner(f"Radim CSV"): 
        converter = DocumentConverter()
        current_date = datetime.now()
        date_string = current_date.strftime('%Y%m%d')
        structured_data = []
        documents = converter.conv_csv(uploaded_file.name)
    
        for doc in documents:
            content = doc.page_content
            metadata = doc.metadata
    
            output_dict = {
                "id": str(uuid4()),
                "chunk": 1,
                "text": content,
                "source": uploaded_file.name,  # Only include the file name
                "date": int(date_string),
                **metadata
            }
            structured_data.append(output_dict)

        json_string = (
            "["
            + ",\n".join(
                json.dumps(d, ensure_ascii=False) for d in structured_data
            )
            + "]"
        )

        return json_string

def dl_json(dokum, json_string):
    napisano = st.info(
                    "Tekstovi su sačuvani u JSON obliku, downloadujte ih na svoj računar"
                )
    file_name = os.path.splitext(dokum.name)[0]
    skinuto = st.download_button(
        "Download JSON",
        data=json_string,
        file_name=f"{file_name}.json",
        mime="application/json",
    )
    return napisano, skinuto

# in myfunc.embeddings.py
class MultiQueryDocumentRetriever:
    """
    Retrieves relevant documents for complex queries by utilizing multi-query techniques
    and Pinecone's hybrid search capabilities.
    
    Attributes:
        
        question (str): Main query for document retrieval.
        namespace (str): Namespace within Pinecone where documents are stored.
        model (str): Model used for generating embeddings for the query and documents.
        temperature (int): Temperature setting for the language model, affecting creativity.
        host (str): Host URL for the Pinecone service.
        tip (str): multiquery or subquery "multi", "sub"
    """

    def __init__(self, question, namespace="zapisnici", model="text-embedding-3-large", host="https://neo-positive-a9w1e6k.svc.apw5-4e34-81fa.pinecone.io", temperature=0, tip="multi"):
        
        self.question = question
        self.namespace = namespace
        self.model = model
        self.host = host
        self.temperature = temperature
        self.tip = tip
        self.log_messages = []
        self.logger = self._init_logger()
        self.llm = ChatOpenAI(model=os.getenv("OPENAI_MODEL"), temperature=self.temperature)
        self.retriever = self._init_retriever()

    def _init_logger(self):
        """Initializes and configures the logger to store messages in an attribute."""
        class ListHandler(logging.Handler):
            def __init__(self, log_messages_list):
                super().__init__()
                self.log_messages_list = log_messages_list

            def emit(self, record):
                log_message = self.format(record)
                self.log_messages_list.append(log_message)

        log_handler = ListHandler(self.log_messages)
        log_handler.setLevel(logging.INFO)
        logger = logging.getLogger("langchain.retrievers.multi_query")
        logger.setLevel(logging.INFO)
        logger.addHandler(log_handler)
        return logger

    def _init_retriever(self):
        """Initializes the document retriever with Pinecone Hybrid Search."""
        api_key = os.environ.get("PINECONE_API_KEY_S")
        pinecone = Pinecone(api_key=api_key, host=self.host)
        index = pinecone.Index(host=self.host)
        
        embeddings = OpenAIEmbeddings(model=self.model)
        bm25_encoder = BM25Encoder()
        bm25_encoder.fit(self.question)

        pinecone_retriever = PineconeHybridSearchRetriever(
            embeddings=embeddings,
            sparse_encoder=bm25_encoder,
            index=index,
            namespace=self.namespace
        )
        if self.tip == "sub":
            our_template = """You are a helpful assistant that generates multiple sub-questions related to an input question. \n
                The goal is to break down the input into a set of sub-problems / sub-questions that can be answers in isolation. \n
                Generate multiple search queries related to: {question} \n
                Output (4 queries): 
                {question}
            """
        else:    
            our_template = """You are an AI language model assistant. Your task is to generate 4 different versions of the given user 
                question to retrieve relevant documents from a vector  database. 
                By generating multiple perspectives on the user question, your goal is to help the user overcome some of the limitations 
                of distance-based similarity search. Provide these alternative questions separated by newlines. 
                Original question: 
                {question}
            """
        our_prompt = PromptTemplate(input_variables=['question'], template=our_template)
        return MultiQueryRetriever.from_llm(retriever=pinecone_retriever, llm=self.llm, prompt=our_prompt)

    def get_relevant_documents(self, custom_question=None):
        """Retrieves documents relevant to the provided or default query."""
        if custom_question is None:
            custom_question = self.question
        
        result = self.retriever.get_relevant_documents(query=custom_question)
        
        return "\n\n".join([doc.page_content for doc in result])


# in myfunc.embeddings.py
class PineconeRetriever:
    """Handles document retrieval using Pinecone's hybrid search capabilities.
    
    Attributes:
        query (str): The search query for retrieving documents.
        namespace (str): The namespace within Pinecone where the documents are stored.
        model (str): The model used for generating embeddings for the query and documents.
    """
    def __init__(self, query, namespace="zapisnici", model="text-embedding-3-large", host="https://neo-positive-a9w1e6k.svc.apw5-4e34-81fa.pinecone.io"):
            self.query = query
            self.namespace = namespace
            self.model = model
            self.host = host  # Ensure the host attribute is correctly set here
            self.pinecone = self._init_pinecone(host)  # Pass the host to the Pinecone initialization
            self.index = self.pinecone.Index(host=self.host)  # Use the host attribute here
            self.embeddings = OpenAIEmbeddings(model=self.model)
            self.bm25_encoder = BM25Encoder()
            self.bm25_encoder.fit(self.query)
            self.retriever = PineconeHybridSearchRetriever(
                embeddings=self.embeddings,
                sparse_encoder=self.bm25_encoder,
                index=self.index,
                namespace=self.namespace
            )

    def _init_pinecone(self, host):
        """Initializes the Pinecone service with the API key and host.
        
        Args:
            host (str): The host URL for the Pinecone service.
        """
        api_key = os.environ.get("PINECONE_API_KEY_S")
        return Pinecone(api_key=api_key, host=host)


    def get_relevant_documents(self):
        """Retrieves documents relevant to the query from Pinecone."""
        return self.retriever.get_relevant_documents(self.query)


# in myfunc.embeddings.py
def prepare_embeddings(chunk_size, chunk_overlap, dokum, sep):
    file_name = "chunks.json"
    semantic = st.radio(
        "Odaberite tip pripreme: ",
        ("Standard", "Heading", "Semantic", "CSV"),
        key="semantic",
        index=None,
        horizontal=True,
        help="Način pripreme JSON fajla za embeding",
    )
    with io.open(dokum.name, "wb") as file:
            file.write(dokum.getbuffer())
    json_string = None
    if semantic == "Standard":
        json_string = standard_chunks(dokum, chunk_size, chunk_overlap, sep)
    elif semantic in ["Heading", "Semantic", "CSV"] and st.button(f"Pripremi {semantic}"):
        with st.spinner(f"Radim {semantic}"):
            if semantic == "Heading":
                json_string = heading_chunks(dokum)
            elif semantic == "Semantic":
                data = read_uploaded_file(dokum)
                # data = pinecone_utility.read_uploaded_file(dokum)
                json_string = semantic_chunks(data, dokum.name)
            elif semantic == "CSV": 
                json_string = csv_chunks(dokum)

    if json_string is not None:
        file_name = os.path.splitext(dokum.name)[0] + ".json"
        # Download button
        dl_button = st.download_button(
            "Download JSON",
            data=json_string,
            file_name=file_name,
            mime="application/json",
        
        )
       

# in myfunc.embeddings.py
def do_embeddings(dokum, tip, api_key, host, index_name, index):
    with st.form(key="my_form_do", clear_on_submit=False):
        err_log = ""
        
        # Now, you can use stored_texts as your texts
        namespace = st.text_input(
            "Unesi naziv namespace-a: ",
            help="Naziv namespace-a je obavezan za kreiranje Pinecone Indeksa",
        )
        submit_b2 = st.form_submit_button(
            label="Submit", help="Pokreće kreiranje Pinecone Indeksa"
        )
        if submit_b2 and dokum and namespace:
            stringio = io.StringIO(dokum.getvalue().decode("utf-8"))

            # Directly load the JSON data from file content
            data = json.load(stringio)

            # Initialize lists outside the loop
            my_list = []
            my_meta = []

            # Process each JSON object in the data
            for item in data:
                # Append the text to my_list
                my_list.append(item['text'])
    
                # Append other data to my_meta
                meta_data = {key: value for key, value in item.items() if key != 'text'}
                my_meta.append(meta_data)
                
            if tip == "hybrid":
               embeddings = OpenAIEmbeddings(model="text-embedding-3-large")
               bm25_encoder = BM25Encoder()
               # fit tf-idf values on your corpus
               bm25_encoder.fit(my_list)

               retriever = PineconeHybridSearchRetriever(
                    embeddings=embeddings,
                    sparse_encoder=bm25_encoder,
                    index=index,
               )
                           
               retriever.add_texts(texts=my_list, metadatas=my_meta, namespace=namespace)
               
            else:
                embed_model = "text-embedding-3-large"
               
                batch_size = 100  # how many embeddings we create and insert at once
                progress_text2 = "Insertovanje u Pinecone je u toku."
                progress_bar2 = st.progress(0.0, text=progress_text2)
                ph2 = st.empty()
            
                for i in tqdm(range(0, len(data), batch_size)):
                    # find end of batch
                    i_end = min(len(data), i + batch_size)
                    meta_batch = data[i:i_end]

                    # get texts to encode
                    ids_batch = [x["id"] for x in meta_batch]
                    texts = [x["text"] for x in meta_batch]
                
                    # create embeddings (try-except added to avoid RateLimitError)
                    try:
                        res = client.embeddings.create(input=texts, model=embed_model)

                    except Exception as e:
                        done = False
                        print(e)
                        while not done:
                            sleep(5)
                            try:
                                res = client.embeddings.create(input=texts, model=embed_model)
                                done = True

                            except:
                                pass

                    embeds = [item.embedding for item in res.data]

                    # Check for [nan] embeddings
              
                    if len(embeds) > 0:
                    
                        to_upsert = list(zip(ids_batch, embeds, meta_batch))
                    else:
                        err_log += f"Greška: {meta_batch}\n"
                    # upsert to Pinecone
                    err_log += f"Upserting {len(to_upsert)} embeddings\n"
                    with open("err_log.txt", "w", encoding="utf-8") as file:
                        file.write(err_log)

                    index.upsert(vectors=to_upsert, namespace=namespace)
                    stodva = len(data)
                    if i_end > i:
                        deo = i_end
                    else:
                        deo = i
                    progress = deo / stodva
                    l = int(deo / stodva * 100)

                    ph2.text(f"Učitano je {deo} od {stodva} linkova što je {l} %")

                    progress_bar2.progress(progress, text=progress_text2)
                    

            # gives stats about index
            st.info("Napunjen Pinecone")

            st.success(f"Sačuvano u Pinecone-u")
            pinecone_stats(index, index_name)


# in myfunc.embeddings.py
class CohereReranker:
    """Reranks documents based on relevance to the query using Cohere's rerank model.

    Attributes:
        query (str): The search query for reranking documents.
    """
    def __init__(self, query):
        self.query = query
        self.client = cohere.Client(os.getenv("COHERE_API_KEY"))

    def rerank(self, documents):
        """Reranks the given documents based on their relevance to the query.

        Args:
            documents (list): A list of document texts to be reranked.

        Returns:
            str: A string containing the top reranked documents concatenated together.
        """
        results = self.client.rerank(query=self.query, documents=documents, top_n=3, model="rerank-multilingual-v2.0")
        return "\n\n".join([result.document['text'] for result in results])


# in myfunc.embeddings.py
class LongContextHandler:
    """Reorders documents to prioritize more relevant content for long contexts.

    This class does not require initialization parameters.
    """
    def reorder(self, documents):
        """Reorders the given documents based on their relevance and context.

        Args:
            documents (list): A list of document texts to be reordered.

        Returns:
            str: A string containing the reordered documents concatenated together.
        """
        reordering = LongContextReorder()
        reordered_docs = reordering.transform_documents(documents)
        return "\n\n".join(reordered_docs)


# in myfunc.embeddings.py
class ContextRetriever:
    """Retrieves and compresses documents to only include the most relevant parts.
    
    Attributes:
        documents (list): A list of documents to be compressed.
        model (str): The model used for compression.
    """
    
    def __init__(self, documents, model=os.getenv("OPENAI_MODEL")):
        self.documents = documents
        self.model = model
        
        response_schemas = [
            ResponseSchema(name="compressed_text", description="The compressed text of the document"),
        ]
        
        output_parser = StructuredOutputParser.from_response_schemas(response_schemas)
        format_instructions = output_parser.get_format_instructions()
        
        self.prompt = PromptTemplate(
            template=mprompts["contextual_compression"],
            input_variables=["documents"],
            partial_variables={"format_instructions": format_instructions},
        )
        self.llm = ChatOpenAI(model=self.model, temperature=0)
        # ovo bar je kako je bilo u primeru
        self.chain = self.prompt | self.llm | output_parser

    def get_compressed_context(self):
        """Compresses the provided documents to include only the most relevant parts."""
        # combine documents into a single string
        compressed_output = self.chain.invoke({"documents": "\n\n".join(self.documents)})

        # just in case...
        if isinstance(compressed_output, dict) and 'compressed_text' in compressed_output:
            return compressed_output['compressed_text']
        else:
            return "Error: Unexpected structure of compressed_output"


# in myfunc.embeddings.py
def rag_tool_answer(prompt, phglob):
    context = " "
    st.session_state.rag_tool = get_structured_decision_from_model(prompt)

    if  st.session_state.rag_tool == "Hybrid":
        processor = HybridQueryProcessor()
        context, scores = processor.process_query_results(prompt)
        # st.info("Score po chunku:")
        # st.write(scores)
        
    # SelfQuery Tool Configuration
    elif  st.session_state.rag_tool == "SelfQuery":
        # Example configuration for SelfQuery
        uvod = mprompts["rag_self_query"]
        prompt = uvod + prompt
        context = SelfQueryPositive(prompt, namespace="selfdemo", index_name="neo-positive")
        
    elif st.session_state.rag_tool == "WebSearchProcess":
         context = web_search_process(prompt)
         
    # Parent Doc Tool Configuration
    elif  st.session_state.rag_tool == "ParentDoc":
        # Define stores
        h_retriever = HybridQueryProcessor(namespace="pos-50", top_k=1)
        h_docstore = HybridQueryProcessor(namespace="pos-2650")

        # Perform a basic search hybrid
        basic_search_result, source_result, chunk = h_retriever.process_query_parent_results(prompt)
        # Perform a search filtered by source (from basic search)
        search_by_source_result = h_docstore.search_by_source(prompt, source_result)
        st.write(f"Osnovni rezultat koji sluzi da nadje prvi: {basic_search_result}")
        st.write(f"Krajnji rezultat koji se vraca: {search_by_source_result}")
        return search_by_source_result

    # Parent Chunks Tool Configuration
    elif  st.session_state.rag_tool == "ParentChunks":
        # Define stores
        h_retriever = HybridQueryProcessor(namespace="zapisnici", top_k=1)
        # Perform a basic search hybrid
        basic_search_result, source_result, chunk = h_retriever.process_query_parent_results(prompt)
        # Perform a search filtered by source and a specific chunk range (both from basic search)
        search_by_chunk_result = h_retriever.search_by_chunk(prompt, source_result, chunk)
        st.write(f"Osnovni rezultat koji sluzi da nadje prvi: {basic_search_result}")
        st.write(f"Krajnji rezultat koji se vraca: {search_by_chunk_result}")
        return search_by_chunk_result

    # Graph Tool Configuration
    elif  st.session_state.rag_tool == "Graph": 
        # Read the graph from the file-like object
        graph = NetworkxEntityGraph.from_gml(st.session_state.graph_file)
        chain = GraphQAChain.from_llm(ChatOpenAI(model=os.getenv("OPENAI_MODEL"), temperature=0), graph=graph, verbose=True)
        rezultat= chain.invoke(prompt)
        context = rezultat['result']

    # Hyde Tool Configuration
    elif  st.session_state.rag_tool == "Hyde":
        # Assuming a processor for Hyde exists
        context = hyde_rag(prompt)

    # MultiQuery Tool Configuration
    elif  st.session_state.rag_tool == "MultiQuery":
        # Initialize the MQDR instance
        retriever_instance = MultiQueryDocumentRetriever(prompt)
        # To get documents relevant to the original question
        context = retriever_instance.get_relevant_documents(prompt)
        output=retriever_instance.log_messages
        generated_queries = output[0].split(": ")[1]
        queries = ast.literal_eval(generated_queries)
        st.info(f"Dodatna pitanja - MultiQuery Alat:")
        for query in queries:
            st.caption(query)

    # RAG Fusion Tool Configuration
    elif  st.session_state.rag_tool == "CohereReranking":
        # Retrieve documents using Pinecone
        pinecone_retriever = PineconeRetriever(prompt)
        docs = pinecone_retriever.get_relevant_documents()
        documents = [doc.page_content for doc in docs]
        
        # Rerank documents using Cohere
        cohere_reranker = CohereReranker(prompt)
        context = cohere_reranker.rerank(documents)
        
    elif  st.session_state.rag_tool == "ContextualCompression":
        # Retrieve documents using Pinecone
        pinecone_retriever = PineconeRetriever(prompt)
        docs = pinecone_retriever.get_relevant_documents()
        documents = [doc.page_content for doc in docs]
       
        # Retrieve and compressed context
        context_retriever = ContextRetriever(documents)
        context = context_retriever.get_compressed_context()
        
    elif  st.session_state.rag_tool == "LongContext":
         # Retrieve documents using Pinecone
        pinecone_retriever = PineconeRetriever(prompt)
        docs = pinecone_retriever.get_relevant_documents()
        documents = [doc.page_content for doc in docs]
        
        # Reorder documents for long context handling
        long_context_handler = LongContextHandler()
        context = long_context_handler.reorder(documents)

    elif  st.session_state.rag_tool == "Calendly":
        # Schedule Calendly meeting
        context = positive_calendly(phglob)

    elif st.session_state.rag_tool == "UploadedDoc":
        # Read text from the uploaded document
        try:
            context = "Text from the document: " + st.session_state.uploaded_doc[0].page_content
        except:
            context = "No text found in the document. Please check if the document is in the correct format."

    elif st.session_state.rag_tool == "WebScrap":
        try:
            context = "Text from the webpage: " + scrape_webpage_text(st.session_state.url_to_scrap)
        except:
            context = "No text found in the webpage. Please check if the URL is correct."
    
    return context
