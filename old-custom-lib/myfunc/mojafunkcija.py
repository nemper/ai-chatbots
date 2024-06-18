# in myfunc.mojafunkcija.py
import base64
import io
import markdown
import html
import os
import pandas as pd
import pdfkit
import PyPDF2
import re
import streamlit as st
import streamlit_authenticator as stauth
import streamlit.components.v1 as components
import tiktoken
import yaml

from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from html2docx import html2docx
from io import StringIO
from openai import APIConnectionError, APIError, RateLimitError
from PIL import Image
from smtplib import SMTP
from yaml.loader import SafeLoader

from langchain.callbacks.base import BaseCallbackHandler
from myfunc.varvars_dicts import work_vars


# in myfunc.mojafunkcija.py
def show_logo():
    """
    Display a logo image in the Streamlit sidebar.

    This function retrieves an image from a specified URL and displays it in the Streamlit sidebar with a defined width.
    The image is displayed using Streamlit's 'image' function.
    """

    with st.sidebar:
        st.image(
            "https://test.georgemposi.com/wp-content/uploads/2023/05/positive-logo-red.jpg",
            width=150,
        )


# in myfunc.mojafunkcija.py
class StreamlitRedirect:
    """
    A class designed to redirect and clean console output for display in Streamlit.

    Attributes:
        output_buffer (StringIO): A buffer to store the cleaned text output.

    This class provides methods to write text to the buffer, cleaning it from escape characters and non-alphabetic symbols, and to retrieve the stored output.
    """

    def __init__(self):
        """
        Initializes the StreamlitRedirect class with an empty StringIO buffer.
        """
        self.output_buffer = StringIO()

    def write(self, text):
        """
        Cleans the input text and writes it to the output buffer.

        Args:
            text (str): The text to be cleaned and written to the buffer.
        """
        cleaned_text = re.sub(r"\x1b[^m]*m|[^a-zA-Z\s]", "", text)
        self.output_buffer.write(cleaned_text + "\n")  # Store the output

    def get_output(self):
        """
        Retrieves the content of the output buffer.

        Returns:
            str: The cleaned text stored in the output buffer.
        """
        return self.output_buffer.getvalue()


# in myfunc.mojafunkcija.py
def tiktoken_len(text):
    """
    Calculates the number of tokens in a given text using the 'tiktoken' tokenizer.

    Args:
        text (str): The text for which the number of tokens is to be calculated.

    Returns:
        int: The number of tokens in the given text.
    """

  
    tokenizer = tiktoken.get_encoding("p50k_base")
    tokens = tokenizer.encode(text, disallowed_special=())
    return len(tokens)


# in myfunc.mojafunkcija.py
def pinecone_stats(index, index_name):
    """
    Displays statistics of a Pinecone index in a Streamlit application.

    Retrieves and formats various statistics about a Pinecone index, and then displays these statistics using Streamlit.

    Args:
        index (Pinecone Index): The Pinecone index object.
        index_name (str): The name of the Pinecone index.
    """

    index_name = index_name
    index_stats_response = index.describe_index_stats()
    index_stats_dict = index_stats_response.to_dict()
    st.subheader("Status indexa:")
    st.write(index_name)
    flat_index_stats_dict = flatten_dict(index_stats_dict)

    # Extract header and content from the index
    header = [key.split("_")[0] for key in flat_index_stats_dict.keys()]
    content = [
        key.split("_")[1] if len(key.split("_")) > 1 else ""
        for key in flat_index_stats_dict.keys()
    ]

    # Create a DataFrame from the extracted data
    df = pd.DataFrame(
        {
            "Header": header,
            "Content": content,
            "Value": list(flat_index_stats_dict.values()),
        }
    )

    # Set the desired number of decimals for float values
    pd.options.display.float_format = "{:.2f}".format

    # Apply formatting to specific columns using DataFrame.style
    styled_df = df.style.apply(
        lambda x: ["font-weight: bold" if i == 0 else "" for i in range(len(x))], axis=1
    ).format({"Value": "{:.0f}"})

    # Display the styled DataFrame as a table using Streamlit
    st.write(styled_df)


# in myfunc.mojafunkcija.py
def flatten_dict(d, parent_key="", sep="_"):
    """
    Flattens a nested dictionary.

    Args:
        d (dict): The dictionary to flatten.
        parent_key (str, optional): A prefix to prepend to keys. Defaults to an empty string.
        sep (str, optional): The separator to use between keys. Defaults to "_".

    Returns:
        dict: The flattened dictionary.
    """

    items = []
    for k, v in d.items():
        new_key = f"{parent_key}{sep}{k}" if parent_key else k
        if isinstance(v, dict):
            items.extend(flatten_dict(v, new_key, sep=sep).items())
        else:
            items.append((new_key, v))
    return dict(items)


# in myfunc.mojafunkcija.py
def def_chunk():
    """
    Defines the size and overlap of chunks in a Streamlit sidebar.

    Allows users to select chunk size and overlap size using sliders in the Streamlit sidebar. The function is specifically designed for setting parameters related to document chunking.

    Returns:
        tuple: A tuple containing the selected chunk size and chunk overlap size.
    """

    with st.sidebar:
        st.info("Odaberite velicinu chunka i overlap")
        chunk_size = st.slider(
            "Set chunk size in characters (50 - 32000)",
            50,
            32000,
            1500,
            step=100,
            help="Velicina chunka odredjuje velicinu indeksiranog dokumenta. Veci chunk obezbedjuje bolji kontekst, dok manji chunk omogucava precizniji odgovor.",
        )
        chunk_overlap = st.slider(
            "Set overlap size in characters (0 - 1000), must be less than the chunk size",
            0,
            1000,
            0,
            step=10,
            help="Velicina overlapa odredjuje velicinu preklapanja sardzaja dokumenta. Veci overlap obezbedjuje bolji prenos konteksta.",
        )
        return chunk_size, chunk_overlap


# in myfunc.mojafunkcija.py
def print_nested_dict_st(d):
    """
    Recursively prints a nested dictionary in a Streamlit application.

    Args:
        d (dict): The nested dictionary to be printed.
    """

    for key, value in d.items():
        if isinstance(value, dict):
            st.write(f"{key}:")
            print_nested_dict_st(value)
        else:
            st.write(f"{key}: {value}")


# in myfunc.mojafunkcija.py
class StreamHandler(BaseCallbackHandler):
    """
    A class for handling stream updates in a Streamlit container.

    Attributes:
        container (Streamlit Container): The container where the updates will be displayed.
        text (str): The current text in the stream.

    Inherits from BaseCallbackHandler to provide specific functionality for handling new tokens in a Streamlit application.
    """

   
    def __init__(self, container):
        """
        Initializes the StreamHandler with a specific Streamlit container.

        Args:
            container (Streamlit Container): The container to display updates in.
        """

        self.container = container
        self.text = ""

    def on_llm_new_token(self, token: str, **kwargs):
        """
        Handles new tokens received from a language model.

        Cleans the token and appends it to the current text displayed in the Streamlit container.

        Args:
            token (str): The new token received from the language model.
        """

        token = (
            token.replace('"', "").replace("{", "").replace("}", "").replace("_", " ")
        )
        self.text += token
        self.container.success(self.text)

    def reset_text(self):
        """
        Resets the current text in the stream to an empty string.
        """

        self.text = ""

    def clear_text(self):
        """
        Clears the text displayed in the Streamlit container.
        """

        self.container.empty()


# in myfunc.mojafunkcija.py
def open_file(filepath):
    """
    Opens and reads the content of a file.

    Args:
        filepath (str): The path to the file to be opened.

    Returns:
        str: The content of the file.
    """

    with open(filepath, "r", encoding="utf-8") as infile:
        sadrzaj = infile.read()
        infile.close()
        return sadrzaj


# in myfunc.mojafunkcija.py
def st_style():
    """
    Applies custom CSS styling to hide certain Streamlit elements.

    This function hides the main menu and the footer in a Streamlit application by injecting CSS into the Streamlit markdown.
    """

    hide_streamlit_style = """
                <style>
                MainMenu {visibility: hidden;}
                footer {visibility: hidden;}
                </style>
                """
    st.markdown(hide_streamlit_style, unsafe_allow_html=True)


# in myfunc.mojafunkcija.py
def positive_login(main, verzija):
    """
    Manages user authentication for a Streamlit application using a YAML configuration.

    Args:
        main (function): The main function of the Streamlit app to run after successful login.
        verzija (str): The version of the application.

    Returns:
        tuple: A tuple containing the user's name, authentication status, and email.
    """

    with open("config.yaml") as file:
        config = yaml.load(file, Loader=SafeLoader)

        authenticator = stauth.Authenticate(
            config["credentials"],
            config["cookie"]["name"],
            config["cookie"]["key"],
            config["cookie"]["expiry_days"],
            config["preauthorized"],
        )

        name, authentication_status, username = authenticator.login(
            "Login to Positive Apps", "main"
        )

        # Get the email based on the name variable
        email = config["credentials"]["usernames"][username]["email"]
        access_level = config["credentials"]["usernames"][username]["access_level"]
        st.session_state["name"] = name
        st.session_state["email"] = email
        st.session_state["access_level"] = access_level

    if st.session_state["authentication_status"]:
        with st.sidebar:
            st.caption(f"Ver 1.0.6")
            authenticator.logout("Logout", "main", key="unique_key")
        # if login success run the program
        main()
    elif st.session_state["authentication_status"] is False:
        st.error("Username/password is incorrect")
    elif st.session_state["authentication_status"] is None:
        st.warning("Please enter your username and password")

    return name, authentication_status, email


# in myfunc.mojafunkcija.py
def init_cond_llm(i=None):
    """
    Initializes conditions for a language model in a Streamlit sidebar.

    Allows users to select a language model and set its temperature using Streamlit widgets.

    Args:
        i (int, optional): An identifier for the widgets, useful when multiple instances are used.

    Returns:
        tuple: A tuple containing the selected model and temperature.
    """


    with st.sidebar:
        st.info("Odaberite Model i temperaturu")
        model = st.selectbox(
            "Odaberite model",
            ("gpt-3.5-turbo", "gpt-3.5-turbo-16k", work_vars["names"]["openai_model"]),
            key="model_key" if i is None else f"model_key{i}",
            help="Modeli se razlikuju po kvalitetu, brzini i ceni upotrebe.",
        )
        temp = st.slider(
            "Set temperature (0=strict, 1=creative)",
            0.0,
            2.0,
            step=0.1,
            key="temp_key" if i is None else f"temp_key{i}",
            help="Temperatura utice na kreativnost modela. Sto je veca temperatura, model je kreativniji, ali i manje pouzdan.",
        )
    return model, temp


# in myfunc.mojafunkcija.py
def greska(e):
    """
    Handles exceptions in a Streamlit application by displaying appropriate warning messages in Serbian.

    Args:
        e (Exception): The exception that occurred.
    """

    if "maximum context length" in str(e):
        st.warning(
            f"Nisam u mogucnosti za zavrsim tekst. Pokusajte sa modelom koji ima veci kontekst.")
    elif "Rate limit" in str(e):
        st.warning(
            f"Nisam u mogucnosti za zavrsim tekst. Broj zahteva modelu prevazilazi limite, pokusajte ponovo za nekoliko minuta.")
    else:
        st.warning(
            f"Nisam u mogucnosti za zavrsim tekst. Pokusajte ponovo za nekoliko minuta. Opis greske je:\n {e}")


# in myfunc.mojafunkcija.py
def convert_input_to_date(ulazni_datum):
    """
    Converts a given input string to a datetime object.

    Args:
        ulazni_datum (str): The input date string in the format 'dd.mm.yyyy.'.

    Returns:
        datetime: A datetime object representing the input date, or None if the format is invalid.
    """

    try:
        date_obj = datetime.strptime(ulazni_datum, "%d.%m.%Y.")
        return date_obj
    except ValueError:
        print("Invalid date format. Please enter a date in the format 'dd.mm.yyyy.'")
        return None
    

# in myfunc.mojafunkcija.py
def parse_serbian_date(date_string):
    """
    Parses a date string in Serbian and converts it to a datetime object.

    Args:
        date_string (str): The date string in Serbian.

    Returns:
        datetime: A datetime object representing the parsed date.
    """

    serbian_month_names = {
        "januar": "January",
        "februar": "February",
        "mart": "March",
        "april": "April",
        "maj": "May",
        "jun": "June",
        "jul": "July",
        "avgust": "August",
        "septembar": "September",
        "oktobar": "October",
        "novembar": "November",
        "decembar": "December"
    }

    date_string = date_string.lower()

    for serbian_month, english_month in serbian_month_names.items():
        date_string = date_string.replace(serbian_month, english_month)

    return datetime.strptime(date_string.strip(), "%d. %B %Y")


# in myfunc.mojafunkcija.py
def send_email(subject, message, from_addr, to_addr, smtp_server, smtp_port, username, password, image_path=None):
    """
    Sends an email using SMTP protocol.

    Args:
    subject (str): Subject line of the email.
    message (str): Body of the email.
    from_addr (str): Sender's email address.
    to_addr (str): Recipient's email address.
    smtp_server (str): Address of the SMTP server to connect to.
    smtp_port (int): Port number for the SMTP server.
    username (str): Username for the SMTP server authentication.
    password (str): Password for the SMTP server authentication.

    This function creates an email message using the specified subject and
    message, sets up a connection to the specified SMTP server, logs in with
    provided credentials, and sends the email. The connection is securely 
    established using TLS (Transport Layer Security).
    """
    
    msg = MIMEMultipart()
    msg['From'] = from_addr
    msg['To'] = to_addr
    msg['Subject'] = subject

    msg.attach(MIMEText(message, 'plain'))
    # Attach the body text
    
    if image_path is not None: # Attach the image
        with open(image_path, 'rb') as file:
            img = MIMEImage(file.read())
            img.add_header('Content-Disposition', 'attachment; filename=image_path')
            msg.attach(img)

    server = SMTP(smtp_server, smtp_port)
    server.starttls()
    server.login(username, password)
    text = msg.as_string()
    server.sendmail(from_addr, to_addr, text)
    server.quit()


# in myfunc.mojafunkcija.py
def sacuvaj_dokument(content, file_name):
    """
    Saves a markdown content as a text, DOCX, and PDF file, providing options to download each format.

    Args:
    content (str): The markdown content to be saved.
    file_name (str): The base name for the output files.

    This function converts the markdown content into HTML, then into a DOCX document
    and a PDF file. It justifies the paragraphs in the DOCX document. The function 
    also provides Streamlit download buttons for each file format: text, DOCX, and PDF.
    The function assumes an environment where Streamlit, markdown, html2docx, and pdfkit
    libraries are available, and uses UTF-8 encoding for the text file.
    """
    st.info("Čuva dokument")
    options = {
        "encoding": "UTF-8",  # Set the encoding to UTF-8
        "no-outline": None,
        "quiet": "",
    }
    
    html = markdown.markdown(content)
    buf = html2docx(html, title="Content")
    # Creating a document object
    doc = Document(io.BytesIO(buf.getvalue()))
    # Iterate over the paragraphs and set them to justified
    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Creating a byte buffer object
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)  # Rewind the buffer to the beginning

    pdf_data = pdfkit.from_string(html, False, options=options)
    
    # strip extension, add suffix
    file_name = os.path.splitext(file_name)[0] + "_out"
    
    st.download_button(
        "Download as .txt",
        content,
        file_name=f"{file_name}.txt",
        help="Čuvanje dokumenta",
    )
            
    st.download_button(
        label="Download as .docx",
        data=doc_io,
        file_name=f"{file_name}.docx",
        mime="docx",
        help= "Čuvanje dokumenta",
    )
            
    st.download_button(
        label="Download as .pdf",
        data=pdf_data,
        file_name=f"{file_name}.pdf",
        mime="application/octet-stream",
        help= "Čuvanje dokumenta",
    )


# in myfunc.mojafunkcija.py
def initialize_session_state(defaults):
    for key, value in defaults.items():
        if key not in st.session_state:
            if callable(value):
                # ako se dodeljuje npr. funkcija
                st.session_state[key] = value()
            else:
                st.session_state[key] = value


# in myfunc.mojafunkcija.py
def check_openai_errors(main_function):
    try:
        main_function()
    except RateLimitError as e:
        if 'insufficient_quota' in str(e):
            st.warning("Potrošili ste sve tokene, kontaktirajte Positive za dalja uputstva")
            # Additional handling, like notifying the user or logging the error
        else:
            st.warning(f"Greška {str(e)}")
    except APIConnectionError as e:
        # Handle connection error here
        st.warning(f"Ne mogu da se povežem sa OpenAI API-jem: {e} pokušajte malo kasnije.")
    except APIError as e:
        # Handle API error here, e.g. retry or log
        st.warning(f"Greška u API-ju: {e} pokušajte malo kasnije.")
    except Exception as e:
        # Handle other exceptions
        st.warning(f"Greška : {str(e)} pokušajte malo kasnije.")


# in myfunc.mojafunkcija.py
def read_docx(file):
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    text_data = '\n'.join(full_text)
    st.write(text_data)
    return text_data


# in myfunc.mojafunkcija.py
def read_txt(file):
    txt_data = file.getvalue().decode("utf-8")
    with st.expander("Prikaži tekst"):
        st.write(txt_data)
    return 


# in myfunc.mojafunkcija.py
def read_csv(file):
    csv_data = pd.read_csv(file)
    with st.expander("Prikaži CSV podatke"):
        st.write(csv_data)
    csv_content = csv_data.to_string()
    return csv_content


# in myfunc.mojafunkcija.py
def read_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    num_pages = len(pdf_reader.pages)
    text_content = ""

    for page in range(num_pages):
        page_obj = pdf_reader.pages[page]
        text_content += page_obj.extract_text()

    # Remove bullet points and fix space issues
    text_content = text_content.replace("•", "")
    text_content = re.sub(r"(?<=\b\w) (?=\w\b)", "", text_content)
    with st.expander("Prikaži tekst"):
        st.write(text_content)
    return text_content


# in myfunc.mojafunkcija.py
def read_image(file):
    base64_image = base64.b64encode(file.getvalue()).decode('utf-8')
    image_bytes = base64.b64decode(base64_image)
    image = Image.open(io.BytesIO(image_bytes))
    with st.expander("Prikaži sliku"):
        st.image(image, width=150)
    return f"data:image/jpeg;base64,{base64_image}"


# in myfunc.mojafunkcija.py
def read_imgs():
    uploaded_file = st.file_uploader("Choose image", help="Odabir dokumenta")
    if uploaded_file is not None:
        if uploaded_file.name.endswith((".jpg", ".jpeg", ".png", ".webp")):
            # Read the image file and convert it to a string
            image_data = read_image(uploaded_file)
            return image_data, True
        else:
            st.error("❌ Greška! Mora slika!")
            return False, False 
    return False, False


# in myfunc.mojafunkcija.py
def read_txts():
    uploaded_files = st.file_uploader("Choose file(s)", accept_multiple_files=True)
    documents = {}
    if uploaded_files:
        for file in uploaded_files:
            filename = file.name
            if filename.endswith('.txt') or filename.endswith('.js') or filename.endswith('.py') or filename.endswith('.md'):
                documents[filename] = read_txt(file)
            elif filename.endswith('.docx'):
                documents[filename] = read_docx(file)
            elif filename.endswith('.pdf'):
                documents[filename] = read_pdf(file)
            elif filename.name.endswith(".csv"):
                documents[filename] = read_csv(file)
            else:
                st.error("❌ Greška! Mora slika!")
                return False, False
        pairs = []
        for key, value in documents.items():
            pairs.append(f"{key}: \n{value}")

        return '\n\n'.join(pairs), True
    return False, False

def copy_to_clipboard(message):
    sanitized_message = html.escape(message)  # Escape the message to handle special HTML characters
    # Create an HTML button with embedded JavaScript for clipboard functionality and custom CSS
    html_content = f"""
    <html>
    <head>
        <style>
            #copyButton {{
                background-color: #454654;  /* Dark gray background */
                color: #f1f1f1;            /* Our white text color */
                border: none;           /* No border */
                border-radius: 8px;     /* Rounded corners */
                cursor: pointer;        /* Pointer cursor on hover */
                outline: none;          /* No focus outline */
                font-size: 20px;        /* Size */
            }}
            #textArea {{
                opacity: 0; 
                position: absolute; 
                pointer-events: none;
            }}
        </style>
    </head>
    <body>
    <textarea id="textArea">{sanitized_message}</textarea>
    <textarea id="textArea" style="opacity: 0; position: absolute; pointer-events: none;">{sanitized_message}</textarea>
    <button id="copyButton" onclick='copyTextToClipboard()'>📄</button>
    <script>
    function copyTextToClipboard() {{
        var textArea = document.getElementById("textArea");
        var copyButton = document.getElementById("copyButton");
        textArea.style.opacity = 1;
        textArea.select();
        try {{
            var successful = document.execCommand('copy');
            var msg = successful ? '✔️' : '❌';
            copyButton.innerText = msg;
        }} catch (err) {{
            copyButton.innerText = 'Failed!';
        }}
        textArea.style.opacity = 0;
        setTimeout(function() {{ copyButton.innerText = "📄"; }}, 3000);  // Change back after 3 seconds
    }}
    </script>
    </body>
    </html>
    """
    components.html(html_content, height=50)